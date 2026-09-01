"""Build the submission report as DOCX (and PDF if LibreOffice is available).

Run:  python make_report.py

Figures are taken from eda_out/ and scorer_results/, so run `python eda.py`,
`python model.py` and `score.py` first. Backtest numbers are the measured output
of `python model.py` and are restated here as constants so report generation
stays fast and deterministic.
"""
from __future__ import annotations

import shutil
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT_DOCX = Path("report/Freight_Rate_Report.docx")
TEAL = RGBColor(0x06, 0x4A, 0x56)
GREY = RGBColor(0x45, 0x5A, 0x60)

BACKTEST = [
    ("Q2 replica (May-Jun)", "9,571", "0.0126", "0.0158", "1.26%"),
    ("Jun-Jul", "9,567", "0.0131", "0.0164", "1.31%"),
    ("Jul-Aug", "9,538", "0.0163", "0.0201", "1.65%"),
    ("Q3 replica (Aug-Sep)", "9,297", "0.0137", "0.0172", "1.38%"),
    ("Sep-Oct", "9,379", "0.0137", "0.0170", "1.36%"),
    ("Mean", "", "0.0139", "0.0173", "1.39%"),
]


def style(doc):
    n = doc.styles["Normal"]
    n.font.name, n.font.size = "Calibri", Pt(10.5)
    n.paragraph_format.space_after = Pt(7)
    for i, size in ((1, 17), (2, 13), (3, 11.5)):
        h = doc.styles[f"Heading {i}"]
        h.font.name, h.font.size, h.font.bold = "Calibri", Pt(size), True
        h.font.color.rgb = TEAL
        h.paragraph_format.space_before = Pt(15 if i < 3 else 11)
        h.paragraph_format.space_after = Pt(5)
        h.paragraph_format.keep_with_next = True


def para(doc, text, italic=False, size=None, color=None, after=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        bold, rest = (it if isinstance(it, tuple) else (None, it))
        if bold:
            p.add_run(bold).bold = True
        p.add_run(rest)


def table(doc, header, rows, widths=None, bold_last=False):
    t = doc.add_table(rows=1, cols=len(header))
    t.style, t.alignment = "Light Grid Accent 1", WD_TABLE_ALIGNMENT.CENTER
    for c, name in zip(t.rows[0].cells, header):
        c.text = ""
        r = c.paragraphs[0].add_run(name)
        r.bold, r.font.size = True, Pt(9.5)
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for c, v in zip(cells, row):
            c.text = ""
            r = c.paragraphs[0].add_run(str(v))
            r.font.size = Pt(9.5)
            if bold_last and i == len(rows) - 1:
                r.bold = True
    if widths:
        for row in t.rows:
            for c, w in zip(row.cells, widths):
                c.width = Inches(w)
    # Keep the whole table on one page: no row splits, and every row but the
    # last pulls the next one along with it.
    for i, row in enumerate(t.rows):
        pr = row._tr.get_or_add_trPr()
        pr.append(OxmlElement("w:cantSplit"))
        if i < len(t.rows) - 1:
            for c in row.cells:
                for pgh in c.paragraphs:
                    pgh.paragraph_format.keep_with_next = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def figure(doc, path, caption, width=6.4):
    if not Path(path).is_file():
        para(doc, f"[missing figure: {path}]", italic=True, color=GREY)
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.keep_with_next = True
    p = para(doc, caption, italic=True, size=8.5, color=GREY)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.7)
        s.left_margin = s.right_margin = Inches(0.8)
    style(doc)

    h = doc.add_heading("Freight Rate Prediction", level=0)
    h.runs[0].font.color.rgb = TEAL
    para(doc, "Spotter Machine Learning Engineer assessment", italic=True, color=GREY)

    # ---------------------------------------------------------------- summary
    doc.add_heading("Summary", level=1)
    para(doc, "Predicting posted_rate for 12,000 loads in November and December 2025 from "
              "48,000 labelled loads spanning January to October. Backtested over five "
              "two-month-ahead folds, the model averages 1.39% MAPE.")
    para(doc, "Two properties of this dataset decide the outcome, and both are invisible to "
              "standard practice. One feature reproduces the target almost exactly in half the "
              "training months and is worthless in the other half, with the validation window "
              "falling in the worthless half. Separately, the equipment premium ramps into "
              "quarter end, and Q4's ramp is the one quarter that carries no labels. Random "
              "cross-validation rewards getting both of these wrong.")

    # ------------------------------------------------------------ data quality
    doc.add_heading("1. Data quality", level=1)
    para(doc, "Four defects, all repaired on the inference path rather than only in training, "
              "because validation is roughly twice as dirty as training:")
    bullets(doc, [
        ("677 corrupted labels (1.41%). ", "340 scaled up 2.2-5.4x, 337 scaled down "
         "0.17-0.44x. Residuals form a clean core, then an entirely empty gap, then two "
         "symmetric lobes, so the 5-sigma cut point is unambiguous. Removed before training."),
        ("Sign-flipped weights. ", "292 training and 145 validation rows carry negative "
         "weight. Absolute values fall in the normal range and then price like their peers, so "
         "abs() is the correct repair."),
        ("Missing values. ", "weight and market_index only, missing completely at random. "
         "market_index is imputed from the same-day mean rather than a global median, because "
         "it is effectively a daily series: daily means explain 98% of its variance."),
        ("Defect rates double at the boundary. ", "market_index missing goes from 0.8% in "
         "training to 2.0% in validation, as a discrete step rather than a drift, so the "
         "validation rate cannot be extrapolated from a training trend."),
    ])
    figure(doc, "eda_out/01_target_and_outliers.png",
           "Figure 1. The corrupted labels sit past an empty gap, making the cut point unambiguous.")

    # ----------------------------------------------------------------- trap 1
    doc.add_heading("2. The quote_signal trap", level=1)
    para(doc, "quote_signal reproduces the target's rate-per-mile almost exactly in January, "
              "February, March, June and September, and is worthless in the other five months.")
    table(doc, ["Months", "Loads within 2% of actual $/mi", "Correlation with target"],
          [["Jan, Feb, Mar, Jun, Sep", "92-94%", "0.95"],
           ["Apr, May, Jul, Aug, Oct", "7-10%", "~0.00"]], widths=[2.6, 2.4, 1.9])
    para(doc, "Under random cross-validation this looks like a jackpot, because half the "
              "training rows carry the answer. On the real holdout it is worthless.")
    para(doc, "Which regime the validation window sits in is determinable without labels, "
              "because the degraded months have a distinct marginal shape. A KS test of "
              "validation against each training month separates it decisively:")
    table(doc, ["Feature", "Best-matching month", "Next best"],
          [["quote_signal", "August  0.009", "January  0.052"],
           ["market_index", "January  0.039", "October  0.152"],
           ["distance", "0.007", "all months under 0.017"],
           ["weight", "0.008", "all months under 0.027"],
           ["all four coordinates", "0.010-0.024", "no shift"]], widths=[2.0, 2.4, 2.5])
    para(doc, "August is a degraded month, and the match is six times tighter than the "
              "runner-up. The feature is excluded from the model entirely. It cannot be "
              "rescued row-wise either: gating on rows where it agrees with a model prediction "
              "leaves residual sd at 0.0230 against a 0.0231 baseline.")
    figure(doc, "eda_out/02_quote_signal_regime.png",
           "Figure 2. Right panel: validation's quote_signal distribution overlays August, a degraded month.")

    # ----------------------------------------------------------------- trap 2
    doc.add_heading("3. The quarter-end trap", level=1)
    para(doc, "Fitting the model separately per month, distance, weight and coordinates are "
              "stable to within a few percent. Equipment is not: Flatbed and Reefer premiums "
              "jump in March, June and September, and only those months. The cause is that "
              "the equipment premium is flat for the first 60 days of a quarter and then ramps.")
    table(doc, ["", "Day 0-59", "Day 84-92"],
          [["Flatbed premium", "+0.068", "+0.140"],
           ["Reefer premium", "+0.113", "+0.160"],
           ["Dry Van (general ramp)", "0.000", "+0.022"]], widths=[2.6, 2.1, 2.1])
    para(doc, "It repeats identically in Q1, Q2 and Q3, and modelling it cuts residual sigma "
              "from 0.0321 to 0.0238, a 26% reduction.")
    para(doc, "The trap is coverage. Training stops on 31 October, day 30 of Q4, so November "
              "falls entirely before the ramp and December entirely inside it. Half the "
              "holdout is an extrapolation into a region with no training coverage. Unlike "
              "quote_signal this is not leakage, so the response is an explicit "
              "day-of-quarter feature rather than a feature to drop.")
    figure(doc, "eda_out/05_equipment_quarter_end.png",
           "Figure 3. The ramp, and the Q4 coverage gap that is exactly December.")

    # ------------------------------------------------- validation and the split
    doc.add_heading("4. Validation approach and data split", level=1)
    para(doc, "The real task trains through day 30 of a quarter and predicts days 31 to 91 of "
              "that same quarter, two months forward. That exact shape occurs twice inside the "
              "training data, so the split is not a generic time cut but a structural replica "
              "of the task:")
    table(doc, ["Split", "Train", "Holdout", "Quarters covering the ramp"],
          [["Q2 replica", "1 Jan - 30 Apr", "1 May - 30 Jun", "Q1 only"],
           ["Q3 replica", "1 Jan - 31 Jul", "1 Aug - 30 Sep", "Q1, Q2"],
           ["Real task", "1 Jan - 31 Oct", "1 Nov - 31 Dec", "Q1, Q2, Q3"]],
          widths=[1.4, 1.8, 1.8, 2.0])
    para(doc, "Both replicas are slightly harder than the real task, since each has fewer "
              "prior quarters from which to learn the ramp. The Q3 replica additionally lands "
              "on degraded-quote_signal months, matching the validation regime.")
    para(doc, "Where two folds were too thin to choose a hyper-parameter, this was widened to "
              "five rolling-origin folds, each still predicting two months forward.")
    doc.add_heading("Why not random cross-validation", level=3)
    bullets(doc, [
        ("It rewards the trap. ", "Half the training rows carry the answer in quote_signal, "
         "and a random split puts those rows on both sides."),
        ("It hides the extrapolation. ", "Random folds contain December-like rows in training, "
         "so the Q4 ramp gap never appears."),
        ("It measures the wrong error. ", "The task is forward in time; random folds "
         "interpolate."),
    ])
    doc.add_heading("The diagnostic that mattered", level=3)
    para(doc, "Average bias hides the failure that matters. A gradient-boosted model's bias "
              "grows across the holdout, from +0.0022 in week 1 to +0.0322 in week 8 on the Q2 "
              "replica, while the chosen model stays flat at +0.0022 to +0.0040. December is "
              "the far end of the holdout, so bias growth is what to select on.")
    para(doc, "That failure is trees being piecewise constant. Training covers day-of-year 1 "
              "to 304 and the holdout is 305 to 365, so every holdout row falls in the top bin "
              "and inherits the last training days' value. Its cost is not fixed: 3.2% on the "
              "Q2 replica where the market was moving, 0.24% on Q3 where it was flat. Which "
              "case December resembles is not knowable in advance, so the trend is carried by "
              "a parametric term that can extrapolate.")
    figure(doc, "eda_out/06_gbm_extrapolation.png",
           "Figure 4. Left: bias growing into the holdout. Right: the fitted curve flattening past the boundary.")

    # ------------------------------------------------------------------- model
    doc.add_heading("5. The model", level=1)
    para(doc, "Four stages, each fitted on the residual of the one before, on log(rate/mile):")
    bullets(doc, [
        ("Parametric core. ", "Linear splines on log(distance) and weight, equipment dummies, "
         "market_index, a linear day-of-year trend, and a quarter-end ramp interacted with "
         "equipment. The trend is linear deliberately, so it extrapolates."),
        ("City premium. ", "One symmetric premium per city applied at both ends of the lane. "
         "A load pays the full premium at origin and at destination (regression slopes 1.013 "
         "and 1.015), and reversing a lane does not change the price. Empirical-Bayes shrunk, "
         "with kNN on lat/lon for the eight cities that appear only in validation."),
        ("Lane effect. ", "Per-lane, shrunk by volume."),
        ("Gradient-boosted correction. ", "On what remains, with day-of-year withheld."),
    ])
    para(doc, "Stage 4 is worth only about 1.4%. Once the structure is explicit the trees have "
              "little left to find, so this is effectively a well-specified GLM with shrunk "
              "fixed effects. That was not the starting assumption; it is where the replica "
              "splits led.")
    doc.add_heading("Geography and the eight unseen cities", level=3)
    para(doc, "Allentown, Charlotte, Chicago, Jackson, Knoxville, Laredo, Norfolk and San "
              "Diego appear only in validation, touching 12.1% of rows. The city premium is a "
              "smooth latitude gradient, so coordinates interpolate to them and a city-name "
              "encoding cannot. Leave-one-city-out puts kNN on lat/lon at R2 0.889 for a "
              "held-out city, against 0.708 for a global plane.")
    doc.add_heading("Seasonality", level=3)
    para(doc, "The spring produce season is present: market_index averages 1.248 across April "
              "to July against 0.944 for August to October. It needs no explicit modelling, "
              "because market_index is already a feature and its November and December values "
              "are supplied in validation.csv. A retail peak is absent from this data; August "
              "to October is the annual trough. Adding an annual harmonic was tested and "
              "rejected, taking mean MAE from 0.0140 to 0.0177, since extrapolating a cycle "
              "fitted on a partial year is unstable.")

    # ----------------------------------------------------------------- results
    doc.add_heading("6. Results", level=1)
    table(doc, ["Fold", "n test", "MAE (log)", "RMSE (log)", "MAPE"], BACKTEST,
          widths=[2.0, 1.1, 1.3, 1.3, 1.1], bold_last=True)
    para(doc, "Trained on 47,323 rows after removing the corrupted labels.")
    doc.add_heading("Output checks", level=3)
    table(doc, ["Check", "Result"],
          [["Predicted rate distribution", "median $2,072 against $2,031 in training"],
           ["Predicted $/mi", "median 2.177 against 2.145 in training"],
           ["December vs November $/mi", "2.206 vs 2.146, the quarter-end ramp"],
           ["Equipment premium, Nov (pre-ramp)", "Flatbed +7.1%, Reefer +12.0%"],
           ["Equipment premium, Dec (in ramp)", "Flatbed +10.3%, Reefer +13.9%"],
           ["The eight unseen cities", "$/mi median 2.181 against 2.177 for the rest"],
           ["Format", "12,000 unique ids, all positive, no NaN; scorer passes"]],
          widths=[2.6, 4.2])
    para(doc, "The November and December equipment premiums land on the values measured in "
              "section 3, which confirms the quarter-end interaction is genuinely applied to "
              "the unlabelled Q4 ramp rather than merely described.")

    # ---------------------------------------------------------------- december
    doc.add_heading("7. The December chart", level=1)
    figure(doc, "scorer_results/candidate_december.png",
           "Figure 5. Produced by the provided score.py. Lexington to Fort Wayne, 360 miles, "
           "Dry Van, 32,000 lb; only the date changes.")
    para(doc, "The rate rises from $832 on 1 December to $885 on 31 December, up 6.3% across "
              "the month, with a weekly ripple.")
    bullets(doc, [
        ("The rise ", "is the quarter-end ramp. December sits entirely at days 61 to 91 of "
         "Q4, inside the ramp region that Q4 itself never labels."),
        ("The ripple ", "is the market_index weekly cycle, peaking Thursday near 1.03 and "
         "troughing Sunday near 0.83, which at an elasticity of 0.131 is roughly a 1.3% swing."),
    ])
    para(doc, "Two independent checks. The shape was predicted from the exploratory analysis "
              "alone, before any model existed, at $837 to $900. And the 32 real Lexington to "
              "Fort Wayne loads in training span $758 to $974, so the curve sits inside the "
              "observed range for this lane.")
    para(doc, "A flat line here would have meant the time structure was missed.")

    # ------------------------------------------------------------- limitations
    doc.add_heading("8. Limitations", level=1)
    para(doc, "The model sits within roughly 20% of an optimistic noise floor, estimated by "
              "fitting in-sample with saturated lane fixed effects. Remaining error is "
              "concentrated in three places, none of which look reducible with this data:")
    bullets(doc, [
        ("Trend uncertainty, about 1% at the December horizon. ", "Bias alternates sign across "
         "folds, so this is genuine uncertainty about an unobservable future rather than a "
         "correctable bias. Tuning the trend window was tested on five folds and made it worse."),
        ("Unseen lanes. ", "12.2% of validation rows sit on lanes absent from training and "
         "receive no lane effect."),
        ("Unseen cities. ", "About 0.6% residual on the part of a new city's premium that "
         "coordinates cannot recover."),
    ])
    para(doc, "With more time I would quantify December uncertainty by bootstrapping the trend "
              "across folds and reporting an interval on the chart, rather than a single line.")

    para(doc, "Full analysis, including the negative results, is in FINDINGS.md in the "
              "repository.", italic=True, color=GREY)

    OUT_DOCX.parent.mkdir(exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"wrote {OUT_DOCX}")
    return OUT_DOCX


def to_pdf(docx_path: Path):
    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice:
        print("LibreOffice not found; DOCX only. Open it and export a PDF to submit.")
        return
    subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir",
                    str(docx_path.parent), str(docx_path)],
                   check=True, capture_output=True, timeout=300)
    pdf = docx_path.with_suffix(".pdf")
    print(f"wrote {pdf}" if pdf.is_file() else "PDF conversion produced no file")


if __name__ == "__main__":
    to_pdf(build())
